"""
FRENTE 1 · CLÍMAX — Ofícios Word em massa
Módulo 2 — Automação de Documentos

O esqueleto vira realidade: 30 ofícios personalizados a partir de UM
modelo. Não escrevemos 30 documentos — escrevemos 1 lógica.

    python3 scripts/frente1_oficios_word.py

Lê  : saida/servidores_limpo.xlsx   (gerado pela Frente 0)
Usa : modelo/oficio_modelo.docx     (com marcadores {{NOME}}, {{CPF}}...)
Grava: saida/oficios/oficio_<numero>_<nome>.docx

------------------------------------------------------------------------
O ESQUELETO QUE SE REPETE (slide 10):
    ler planilha → PARA CADA linha → preencher modelo → salvar
A única peça que troca entre Word/PDF/Excel é "preencher / gerar saída".
------------------------------------------------------------------------
"""

import unicodedata
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document

RAIZ = Path(__file__).resolve().parent.parent.parent
PLANILHA = RAIZ / "saida" / "modulo2" / "servidores_limpo.xlsx"
MODELO = RAIZ / "modelo" / "oficio_modelo.docx"
PASTA_SAIDA = RAIZ / "saida" / "modulo2" / "oficios"

# Dados fixos do órgão (iguais para todos os ofícios).
CIDADE = "Teresina"

# Meses por extenso para escrever a data de hoje "à moda de ofício".
MESES = [
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
]


def data_por_extenso(d: date) -> str:
    """22 de junho de 2026 — formato usual de um ofício."""
    return f"{d.day} de {MESES[d.month - 1]} de {d.year}"


def nome_de_arquivo(numero: str, nome: str) -> str:
    """Transforma 'Maria Fernanda Rocha' + '3' em 'oficio_03_maria_fernanda_rocha.docx'.

    Tira acentos e espaços para gerar um nome de arquivo seguro em qualquer
    sistema (abstração: o nome bonito é para humanos; o do arquivo, para a máquina).
    """
    sem_acento = (
        unicodedata.normalize("NFKD", nome)
        .encode("ascii", "ignore")
        .decode("ascii")
    )
    base = sem_acento.lower().replace(" ", "_")
    return f"oficio_{int(numero):02d}_{base}.docx"


def substituir_marcadores(doc: Document, dados: dict[str, str]) -> None:
    """Troca cada {{CAMPO}} pelo valor correspondente, em todo o documento.

    DETALHETransforma IMPORTANTE (e armadilha clássica): o Word costuma quebrar um
    texto em vários "runs" internos, então '{{NOME}}' pode estar partido
    em pedaços ('{{NO' + 'ME}}'). Por isso juntamos o texto do parágrafo
    inteiro, fazemos a troca, e reescrevemos o parágrafo num run só.
    """
    for paragrafo in doc.paragraphs:
        texto = paragrafo.text
        if "{{" not in texto:
            continue
        for chave, valor in dados.items():
            texto = texto.replace("{{" + chave + "}}", str(valor))
        # Reescreve o parágrafo: apaga os runs antigos e põe o texto trocado.
        for run in list(paragrafo.runs):
            run.text = ""
        if paragrafo.runs:
            paragrafo.runs[0].text = texto
        else:
            paragrafo.add_run(texto)


def gerar_um_oficio(linha: pd.Series, hoje: str) -> Path:
    """Gera UM ofício a partir de UMA linha da planilha. (a peça que se repete)"""
    dados = {
        "NOME": linha["Nome"],
        "CPF": linha["CPF"],
        "CARGO": linha["Cargo"],
        "LOTACAO": linha["Lotacao"],
        "DATA_ADMISSAO": linha["Data_Admissao"],
        "EMAIL": linha["Email"],
        "NUMERO_OFICIO": linha["Numero_Oficio"],
        "CIDADE": CIDADE,
        "DATA_HOJE": hoje,
        "ASSUNTO": "Ofício de Comunicação",
    }
    doc = Document(MODELO)               # abre o modelo do zero a cada linha
    substituir_marcadores(doc, dados)    # preenche
    destino = PASTA_SAIDA / nome_de_arquivo(linha["Numero_Oficio"], linha["Nome"])
    doc.save(destino)                    # salva um arquivo por servidor
    return destino


def main() -> None:
    if not PLANILHA.exists():
        raise SystemExit(
            f"Não achei {PLANILHA.relative_to(RAIZ)}.\n"
            "Rode antes a Frente 0:  python3 scripts/frente0_ler_e_limpar.py"
        )

    df = pd.read_excel(PLANILHA, dtype=str)
    PASTA_SAIDA.mkdir(parents=True, exist_ok=True)
    hoje = data_por_extenso(date.today())

    # O LOOP: para cada linha, preencher e salvar. Esse é o coração da aula.
    total = 0
    for _, linha in df.iterrows():
        destino = gerar_um_oficio(linha, hoje)
        print(f"  [ok] {destino.name}")
        total += 1

    print(f"\n[ok] {total} ofícios gerados em: {PASTA_SAIDA.relative_to(RAIZ)}")


if __name__ == "__main__":
    main()
